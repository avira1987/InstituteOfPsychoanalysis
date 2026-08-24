#!/usr/bin/env python3
"""تولید فایل Word راهنمای تست فرم یکپارچه و پنل عمومی فرایند — راست‌چین با فونت ایران‌سنس.

اجرا از ریشهٔ ریپو:
  python scripts/generate_unified_form_generic_panel_test_guide_docx.py

خروجی:
  docs/راهنمای_تست_فرم_یکپارچه_و_پنل_عمومی_فرایند.docx
"""
from __future__ import annotations

import re
import shutil
import sys
import uuid
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

if sys.platform == "win32" and hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_REG = ROOT / "app" / "assets" / "fonts" / "IRANSans-Regular.ttf"
FONT_BOLD = ROOT / "app" / "assets" / "fonts" / "IRANSans-Bold.ttf"
OUT_DOCX = ROOT / "docs" / "راهنمای_تست_فرم_یکپارچه_و_پنل_عمومی_فرایند.docx"

SOURCES = [
    ROOT / "docs" / "unified_form_generic_panel_test_guide_plain_fa.md",
    ROOT / "docs" / "unified_form_generic_panel_role_daily_tasks_plain_fa.md",
]

# نام فونت در فایل TTF: IRANSans (ایران‌سنس)
FONT_NAME = "IRANSans"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

ET.register_namespace("w", W_NS)
ET.register_namespace("r", R_NS)


def _ensure_fonts() -> None:
    if FONT_REG.is_file():
        return
    import urllib.request

    FONT_REG.parent.mkdir(parents=True, exist_ok=True)
    url = (
        "https://raw.githubusercontent.com/RooyeKhat-Media/iGap-Plus/master/"
        "android/app/src/main/assets/fonts/IRANSans.ttf"
    )
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    FONT_REG.write_bytes(urllib.request.urlopen(req, timeout=60).read())


def _apply_run_font(run, *, size: float = 11, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:rFonts", "w:rtl", "w:cs"):
        existing = r_pr.find(qn(tag))
        if existing is not None:
            r_pr.remove(existing)
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    r_pr.insert(0, r_fonts)
    rtl = OxmlElement("w:rtl")
    r_pr.append(rtl)


def _set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:jc", "w:bidi"):
        existing = p_pr.find(qn(tag))
        if existing is not None:
            p_pr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    p_pr.append(jc)
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)


def _set_cell_rtl(cell) -> None:
    for paragraph in cell.paragraphs:
        _set_rtl(paragraph)
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:tcBidi")) is None:
        tc_pr.append(OxmlElement("w:tcBidi"))


def _add_formatted_runs(paragraph, text: str, *, size: float = 11, bold_default: bool = False) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            _apply_run_font(run, size=size, bold=bold_default)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _apply_run_font(run, size=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(size - 0.5)
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        elif token.startswith("["):
            label = re.match(r"\[([^\]]+)\]", token)
            if label:
                run = paragraph.add_run(label.group(1))
                _apply_run_font(run, size=size)
                run.italic = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _apply_run_font(run, size=size, bold=bold_default)


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    return bool(s) and set(s.replace("|", "").replace(":", "").replace("-", "").strip()) == set()


def _parse_table_row(line: str) -> list[str]:
    return [p.strip() for p in line.strip().strip("|").split("|")]


def _set_rtl_style(style) -> None:
    style.font.name = FONT_NAME
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = style.element.get_or_add_pPr()
    for tag in ("w:jc", "w:bidi"):
        existing = p_pr.find(qn(tag))
        if existing is not None:
            p_pr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    p_pr.append(jc)
    p_pr.append(OxmlElement("w:bidi"))


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    sect_pr = section._sectPr
    if sect_pr.find(qn("w:bidi")) is None:
        sect_pr.append(OxmlElement("w:bidi"))

    settings = doc.settings.element
    if settings.find(qn("w:bidiVisual")) is None:
        settings.append(OxmlElement("w:bidiVisual"))
    theme_lang = settings.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_lang)
    theme_lang.set(qn("w:bidi"), "fa-IR")
    theme_lang.set(qn("w:val"), "en-US")

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    _set_rtl_style(normal)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            style.font.bold = True
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            _set_rtl_style(style)

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            _set_rtl_style(doc.styles[style_name])


def _add_paragraph(doc: Document, text: str = "", *, style: str | None = None, size: float = 11):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _set_rtl(p)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    if text:
        _add_formatted_runs(p, text, size=size)
    return p


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    if tbl_pr.find(qn("w:bidiVisual")) is None:
        tbl_pr.append(OxmlElement("w:bidiVisual"))
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            value = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            _set_rtl(p)
            _add_formatted_runs(p, value, size=9.5, bold_default=(r_idx == 0))
            _set_cell_rtl(cell)
    doc.add_paragraph()


def _add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    _set_rtl(p)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def markdown_to_docx(doc: Document, md_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            _add_paragraph(doc, stripped[2:].strip(), style="Heading 1")
            i += 1
            continue

        if stripped.startswith("## "):
            _add_paragraph(doc, stripped[3:].strip(), style="Heading 2")
            i += 1
            continue

        if stripped.startswith("### "):
            _add_paragraph(doc, stripped[4:].strip(), style="Heading 3")
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            table_rows: list[list[str]] = [_parse_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(_parse_table_row(lines[i]))
                i += 1
            _add_table(doc, table_rows)
            continue

        if re.match(r"^[-*] \[ \]", stripped):
            text = re.sub(r"^[-*] \[ \]\s*", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _set_rtl(p)
            _add_formatted_runs(p, "☐ " + text, size=10.5)
            i += 1
            continue

        if re.match(r"^[-*] ", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _set_rtl(p)
            _add_formatted_runs(p, text, size=10.5)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _set_rtl(p)
            _add_formatted_runs(p, text, size=10.5)
            i += 1
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        _add_paragraph(doc, stripped, size=11)
        i += 1


def _embed_fonts(docx_path: Path) -> None:
    """جاسازی فونت ایران‌سنس داخل فایل Word برای نمایش بدون نصب فونت."""
    fonts: list[tuple[str, Path]] = [(FONT_NAME, FONT_REG)]
    if FONT_BOLD.is_file():
        fonts.append((f"{FONT_NAME} Bold", FONT_BOLD))

    temp_dir = docx_path.with_suffix(".embed_tmp")
    if temp_dir.exists():
        shutil.rmtree(temp_dir)
    temp_dir.mkdir()

    with zipfile.ZipFile(docx_path, "r") as zin:
        zin.extractall(temp_dir)

    word_dir = temp_dir / "word"
    fonts_dir = word_dir / "fonts"
    fonts_dir.mkdir(exist_ok=True)
    rels_dir = word_dir / "_rels"
    rels_dir.mkdir(exist_ok=True)

    font_table_path = word_dir / "fontTable.xml"
    if font_table_path.exists():
        font_root = ET.parse(font_table_path).getroot()
    else:
        font_root = ET.Element(f"{{{W_NS}}}fonts")

    rels_path = rels_dir / "fontTable.xml.rels"
    if rels_path.exists():
        rels_root = ET.parse(rels_path).getroot()
    else:
        rels_root = ET.Element(f"{{{REL_NS}}}Relationships")

    ct_path = temp_dir / "[Content_Types].xml"
    ct_root = ET.parse(ct_path).getroot()

    existing_rids = []
    for rel in rels_root.findall(f"{{{REL_NS}}}Relationship"):
        existing_rids.append(rel.get("Id", ""))

    def next_rid() -> str:
        n = 1
        while True:
            candidate = f"rId{n}"
            if candidate not in existing_rids:
                existing_rids.append(candidate)
                return candidate
            n += 1

    font_key = "{" + str(uuid.uuid4()).upper() + "}"

    for font_label, font_file in fonts:
        target_name = font_file.name
        zip_font_path = fonts_dir / target_name
        shutil.copy2(font_file, zip_font_path)

        part_name = f"/word/fonts/{target_name}"
        override_tag = f"{{{CT_NS}}}Override"
        if not any(el.get("PartName") == part_name for el in ct_root.findall(override_tag)):
            override = ET.Element(override_tag)
            override.set("PartName", part_name)
            override.set("ContentType", "application/x-fontdata")
            ct_root.append(override)

        rid = next_rid()
        rel = ET.Element(f"{{{REL_NS}}}Relationship")
        rel.set("Id", rid)
        rel.set(
            "Type",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font",
        )
        rel.set("Target", f"fonts/{target_name}")
        rels_root.append(rel)

        font_el = ET.SubElement(font_root, f"{{{W_NS}}}font")
        font_el.set(f"{{{W_NS}}}name", FONT_NAME if "Bold" not in font_label else FONT_NAME)
        charset = ET.SubElement(font_el, f"{{{W_NS}}}charset")
        charset.set(f"{{{W_NS}}}val", "B2")
        family = ET.SubElement(font_el, f"{{{W_NS}}}family")
        family.set(f"{{{W_NS}}}val", "auto")
        pitch = ET.SubElement(font_el, f"{{{W_NS}}}pitch")
        pitch.set(f"{{{W_NS}}}val", "variable")
        embed_tag = (
            "embedBold" if "Bold" in font_label or font_file == FONT_BOLD else "embedRegular"
        )
        embed = ET.SubElement(font_el, f"{{{W_NS}}}{embed_tag}")
        embed.set(f"{{{R_NS}}}id", rid)
        embed.set(f"{{{W_NS}}}fontKey", font_key)

    ET.ElementTree(font_root).write(font_table_path, encoding="UTF-8", xml_declaration=True)
    ET.ElementTree(rels_root).write(rels_path, encoding="UTF-8", xml_declaration=True)
    ET.ElementTree(ct_root).write(ct_path, encoding="UTF-8", xml_declaration=True)

    doc_rels_path = rels_dir / "document.xml.rels"
    if doc_rels_path.exists():
        doc_rels = ET.parse(doc_rels_path).getroot()
        font_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable"
        if not any(rel.get("Type") == font_rel_type for rel in doc_rels.findall(f"{{{REL_NS}}}Relationship")):
            rid = next_rid()
            rel = ET.Element(f"{{{REL_NS}}}Relationship")
            rel.set("Id", rid)
            rel.set("Type", font_rel_type)
            rel.set("Target", "fontTable.xml")
            doc_rels.append(rel)
            ET.ElementTree(doc_rels).write(doc_rels_path, encoding="UTF-8", xml_declaration=True)

    rebuilt = docx_path.with_suffix(".rebuilt.docx")
    if rebuilt.exists():
        rebuilt.unlink()
    with zipfile.ZipFile(rebuilt, "w", zipfile.ZIP_DEFLATED) as zout:
        for file_path in sorted(temp_dir.rglob("*")):
            if file_path.is_file():
                zout.write(file_path, file_path.relative_to(temp_dir).as_posix())
    shutil.rmtree(temp_dir)
    docx_path.unlink()
    rebuilt.rename(docx_path)


def main() -> None:
    _ensure_fonts()
    if not FONT_REG.is_file():
        raise FileNotFoundError(f"فونت ایران‌سنس یافت نشد: {FONT_REG}")

    for src in SOURCES:
        if not src.is_file():
            raise FileNotFoundError(f"فایل منبع یافت نشد: {src}")

    doc = Document()
    _configure_document(doc)

    cover = doc.add_paragraph()
    _set_rtl(cover)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("راهنمای تست فرم یکپارچه و پنل عمومی فرایند\nمرکز انستیتو روانکاوی تهران")
    _apply_run_font(run, size=18, bold=True)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    _set_rtl(sub)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("شامل: چک‌لیست تست موج ۱ و ۲، و وظایف روزانه نقش‌ها")
    _apply_run_font(sub_run, size=12)

    doc.add_page_break()

    for idx, src in enumerate(SOURCES):
        if idx > 0:
            doc.add_page_break()
        markdown_to_docx(doc, src)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    out_path = OUT_DOCX
    try:
        doc.save(str(out_path))
    except PermissionError:
        out_path = OUT_DOCX.with_name(OUT_DOCX.stem + "_جدید.docx")
        doc.save(str(out_path))
        print(f"توجه: فایل قبلی باز بود؛ نسخه جدید ذخیره شد.")
    _embed_fonts(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()
